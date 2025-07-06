#README
# This script reads data from an Excel file and fills a Word document template with the data.
# Before running this script, ensure you have the required libraries installed:
# pip install openpyxl python-docx pandas

# Import necessary libraries
import os
from docx import Document
from openpyxl import load_workbook
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt
import warnings
warnings.filterwarnings("ignore", category=UserWarning)
import pandas as pd
import json

#Define functions to read and format data from Excel, fill Word tables, and replace text in Word documents.
def format_value(cell):
    value = cell.value
    if value is None:
        return ''
    number_format = cell.number_format
    if '0%' in number_format or '%' in number_format:
        return f"{value * 100:.2f}%" #Format as percentage with two decimal places
    return str(value)
def read_excel_data(excel_path, sheet_name, start_cells=1):
    workbook = load_workbook(excel_path, read_only=True, data_only=True) #Open the workbook in data-only mode which means it will not evaluate formulas, just return the values.
    sheet = workbook[sheet_name] 

    if sheet_name == '表1.基本資料':
        data = {
            'A': [],
            'C': [],
        } #Initialize a dictionary for company name (A column) and address (C column)
        row = 18
        while True:
            cell_a = f'A{row}'
            cell_c = f'C{row}'
            value_a = sheet[cell_a].value
            value_c = sheet[cell_c].value
            if value_a is None or str(value_a).strip() == '':
                break
            data['A'].append(value_a)
            data['C'].append(value_c)
            row += 1 #Read the columns A and C starting from row 13 until an empty cell is found.
    elif sheet_name == '表2.排放源鑑別':
        data = {
            'B': [],
            'C': [],
            'E': [],
            'K': [],
            'I':[],
            'C_category3': [],
            'C_category5': [],
            'C_category6': [],
            'C_category7': [],
            'C_category8': [],
            'C_category10': [],
            'C_category11': [],
            'C_category13': [],
            'C_category14': [],
            'C_category15': [],
            'K_category1':[],
            'C_category1':[],
            'others': []
        }
        row = 4
        while True:
            cell_b = f'B{row}'
            cell_c = f'C{row}'
            cell_e = f'E{row}'
            cell_k = f'K{row}'
            cell_i = f'I{row}'
            value_b = sheet[cell_b].value
            value_c = sheet[cell_c].value
            value_e = sheet[cell_e].value
            value_k = sheet[cell_k].value
            value_i = sheet[cell_i].value #There's nothing in I column?????
            if value_b is None or str(value_c).strip() == '':
                break
            data['B'].append(value_b)
            data['C'].append(value_c)
            data['E'].append(value_e)
            data['K'].append(value_k)
            data['I'].append(value_i)
            data['others'].append('請輸入文字')
            if value_e == '範疇1':
                data['K_category1'].append(value_k)
                data['C_category1'].append(value_c)
            elif value_e == '類別3':
                data['C_category3'].append(value_c)
            elif value_e == '類別5':
                data['C_category5'].append(value_c)
            elif value_e == '類別6':
                data['C_category6'].append(value_c)
            elif value_e == '類別7':
                data['C_category7'].append(value_c)
            elif value_e == '類別8':
                data['C_category8'].append(value_c)
            elif value_e == '類別10':
                data['C_category10'].append(value_c)
            elif value_e == '類別11':
                data['C_category11'].append(value_c)
            elif value_e == '類別13':
                data['C_category13'].append(value_c)
            elif value_e == '類別14':
                data['C_category14'].append(value_c)
            elif value_e == '類別15':
                data['C_category15'].append(value_c)
            row += 1
    elif sheet_name == '表3.活動數據':
        data = {
            'C': [],
            'I':[],
            'others': []
            }
        row = 4
        while True:
            cell_c = f'C{row}'
            cell_i = f'I{row}'
            value_c = sheet[cell_c].value
            value_i = sheet[cell_i].value
            if value_c is None or str(value_c).strip() == '':
                break
            data['C'].append(value_c)
            data['I'].append(value_i)
            data['others'].append('請輸入文字')
            row += 1

    elif sheet_name == '表8.不確定分析':
        data = {
            'B': [],
            'C': [],
            'D': [],
            'E': [],
            'F': [],
            'G': [],
            'H': [],
            'I': [],
            'J': [],
            'K': [],
            'L': [],
            'M': [],
        }
        row = 4
        while True:
            cell_B = f'B{row}'
            cell_C = f'C{row}'
            cell_D = f'D{row}'
            cell_E = f'E{row}'
            cell_F = f'F{row}'
            cell_G = f'G{row}'
            cell_H = f'H{row}'
            cell_I = f'I{row}'
            cell_J = f'J{row}'
            value_B = sheet[cell_B].value
            value_C = sheet[cell_C].value
            value_D = sheet[cell_D].value
            value_E = sheet[cell_E].value
            value_F = sheet[cell_F].value
            value_G = sheet[cell_G].value
            value_H = sheet[cell_H].value
            value_I = sheet[cell_I].value
            value_J = sheet[cell_J].value
            if value_B is None or str(value_C).strip() == '':
                break
            data['B'].append(value_B)
            data['C'].append(value_C)
            data['D'].append(value_D)
            data['E'].append(value_E)
            data['F'].append(value_F)
            data['G'].append(value_G)
            data['H'].append(value_H)
            data['I'].append(value_I)
            data['J'].append(value_J)
            row += 1
    workbook.close()
    return data

def read_excel_data_pandas(excel_path, sheet_name):
    df = pd.read_excel(excel_path, sheet_name=sheet_name, header=2) 
    data = {}
    if sheet_name == '表5.排放係數':
        df = df.dropna(subset=["排放類別"])
        gases = ["CO2", "CH4", "N2O", "HFCS", "PFCS", "SF6", "NF3"]
        transformed_rows = []
        for _, row in df.iterrows():
            has_valid_gas = any(not pd.isna(row.get(gas)) and row.get(gas) != '' for gas in gases)
            if not has_valid_gas:
                continue
            for gas in gases:
                value = row.get(gas)
                if pd.isna(value) or value == '':
                    continue

                try:
                    formatted_value = f"{float(value):.10f}"
                except:
                    formatted_value = str(value)

                transformed_rows.append({
                    "範疇或類別": row["排放類別"],
                    "排放源": row["排放源"],
                    "係數來源": row["係數來源"],
                    "係數名稱": row["係數名稱"],
                    "氣體": gas,
                    "溫室氣體排放係數": formatted_value,
                    "單位": row["單位"]
                })
        final_df = pd.DataFrame(transformed_rows)

        final_df = final_df.fillna("")

        data = {
            '範疇或類別': final_df['範疇或類別'].tolist(),
            '排放源': final_df['排放源'].tolist(),
            '係數來源': final_df['係數來源'].tolist(),
            '係數名稱': final_df['係數名稱'].tolist(),
            '氣體': final_df['氣體'].tolist(),
            '溫室氣體排放係數': final_df['溫室氣體排放係數'].tolist(),
            '單位': final_df['單位'].tolist()
        }
    
    return data

def read_excel_cell(excel_path, sheet_name, cell):
    try:
        workbook = load_workbook(excel_path, data_only=True)
        sheet = workbook[sheet_name]
        value = sheet[cell].value
        workbook.close()
        return str(value) if value is not None else ''
    except Exception as e:
        print(f"讀取儲存格 {cell} 失敗: {str(e)}")
        return ''

def read_excel_cells(excel_path, sheet_name, cells):
    try:
        workbook = load_workbook(excel_path, read_only=False, data_only=True)
        sheet = workbook[sheet_name]
        values = {cell: format_value(sheet[cell]) for cell in cells}
        workbook.close()
        return values
    except Exception as e:
        print(f"批量讀取儲存格失敗: {str(e)}")
        return {cell: '' for cell in cells}

def add_table_row(table):
    tr = OxmlElement('w:tr')
    for _ in range(len(table.columns)):
        tc = OxmlElement('w:tc')
        tc.append(OxmlElement('w:p'))
        tr.append(tc)
    table._tbl.append(tr)

def fill_word_table(word_path, output_path, table_index, excel_data, cell_mapping, start_row=0):
    doc = Document(word_path)
    table = doc.tables[table_index]

    table.autofit = False
    table.allow_autofit = False

    column_widths = [2000, 2000, 2000, 2000]
    for col_idx, width in enumerate(column_widths[:len(table.columns)]):
        for row in table.rows:
            cell = row.cells[col_idx]
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(width))
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)

    max_data_len = max(len(excel_data.get(key, [])) for key in cell_mapping.keys())

    required_rows = start_row + max_data_len
    while len(table.rows) < required_rows:
        add_table_row(table)

    for key, (row_offset, col) in cell_mapping.items():
        for i, value in enumerate(excel_data.get(key, [])):
            cell = table.cell(start_row + i, col)

            # Remove all existing paragraphs
            for para in cell.paragraphs:
                p = para._element
                p.getparent().remove(p)
        
            # Add new clean paragraph and run
            new_para = cell.add_paragraph()
            run = new_para.add_run(str(value).strip())
            run.font.size = Pt(12)
            run.font.name = 'Times New Roman'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            no_wrap = tcPr.find(qn('w:noWrap'))
            if no_wrap is not None:
                tcPr.remove(no_wrap)

    doc.save(output_path)

def replace_texts_in_word(word_path, output_path, replacements):
    doc = Document(word_path)

    for paragraph in doc.paragraphs:
        original_text = paragraph.text
        modified = False
        new_text = original_text
        for old_text, new_text_value in replacements:
            if old_text in new_text:
                new_text = new_text.replace(old_text, new_text_value)
                modified = True
        if modified:
            for run in paragraph.runs:
                run.clear()
            paragraph.clear()
            run = paragraph.add_run(new_text)
            run.font.size = Pt(12)
            run.font.name = 'Times New Roman'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                original_text = cell.text
                new_text = original_text
                modified = False
                for old_text, new_value in replacements:
                    if old_text in new_text:
                        new_text = new_text.replace(old_text, new_value)
                        modified = True
                if modified:
                    # FULLY remove all paragraphs in the cell
                    for para in cell.paragraphs:
                        p = para._element
                        p.getparent().remove(p)

                    # Add clean paragraph
                    paragraph = cell.add_paragraph()
                    run = paragraph.add_run(new_text.strip())
                    run.font.size = Pt(12)
                    run.font.name = 'Times New Roman'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    doc.save(output_path)

def merge_cells_in_table_25(word_path, output_path, table_index=25):
    doc = Document(word_path)
    table = doc.tables[table_index]

    previous_source = None
    merge_start = None

    for row_idx in range(1, len(table.rows)):  # skip header
        current_source = table.cell(row_idx, 1).text.strip()

        if current_source == previous_source:
            continue
        else:
            if merge_start is not None and row_idx - merge_start > 1:
                for col in [0, 1, 2, 3]:  # columns to merge
                    cell_to_merge = table.cell(merge_start, col)
                    for merge_row in range(merge_start + 1, row_idx):
                        # Clear text before merging
                        table.cell(merge_row, col).text = ""
                        cell_to_merge.merge(table.cell(merge_row, col))

            previous_source = current_source
            merge_start = row_idx

    # Handle the last group
    if merge_start is not None and len(table.rows) - merge_start > 1:
        for col in [0, 1, 2, 3]:
            cell_to_merge = table.cell(merge_start, col)
            for merge_row in range(merge_start + 1, len(table.rows)):
                table.cell(merge_row, col).text = ""
                cell_to_merge.merge(table.cell(merge_row, col))

    doc.save(output_path)

def insert_if_empty_tables(word_path, output_path, table_indices):
    doc = Document(word_path)

    for table_index in table_indices:
        table = doc.tables[table_index]

        # Check if all cells in data rows (excluding header) are empty
        is_data_empty = True
        for row in table.rows[1:]:  # assuming row 0 is the header
            if any(cell.text.strip() for cell in row.cells):
                is_data_empty = False
                break

        if is_data_empty:
            # Make sure the table has at least two rows
            while len(table.rows) < 2:
                table.add_row()
            target_cell = table.rows[1].cells[0]  # Insert into the first column
            target_cell.text = "無"
            run = target_cell.paragraphs[0].runs[0]
            run.font.size = Pt(12)
            run.font.name = 'Times New Roman'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    doc.save(output_path)

#Edit the FILEPATH by unhiding the def main() function below.
def main_with_inputs(excel_path, word_path, output_folder, output_file_name):
    output_path = os.path.join(output_folder, output_file_name)

    #excel_path = r"D:\user\Desktop\0611_Testing\CSM_input.xlsx"
    #word_path = r"D:\user\Desktop\RB_Fionna\溫室氣體盤查報告書_模板_0610.docx"
    #output_file_name = "output_SDI.docx"
    #output_path = rf"D:\user\Desktop\{output_file_name}"

    start_cells = ['A13', 'C13', 'E13']

    cell_mapping_table0 = {
        'A': (0, 0),
        'C': (0, 1),
    }
    cell_mapping_table1 = {
        'K_category1': (0, 0),
        'C_category1': (0, 1),
    }
    cell_mapping_table3 = {
        'C_category3': (0, 0),
    }
    cell_mapping_table5 = {
        'C_category5': (0, 0),
    }

    cell_mapping_table6 = {
        'C_category6': (0, 0),
    }

    cell_mapping_table7 = {
        'C_category7': (0, 0),
    }

    cell_mapping_table8 = {
        'C_category8': (0, 0),
    }

    cell_mapping_table10 = {
        'C_category10': (0, 0),
    }

    cell_mapping_table11 = {
        'C_category11': (0, 0),
    }
    cell_mapping_table13 = {
        'C_category13': (0, 0),
    }
    cell_mapping_table14 = {
        'C_category14': (0, 0),
    }
    cell_mapping_table15 = {
        'C_category15': (0, 0),
    }
    cell_mapping_table16 = {
        'E': (0, 0),
        'K': (0, 1),
        'B': (0, 2),
        'C': (0, 3)
    }

    cell_mapping_table23 = {
        'C': (0, 0),
        'I': (0, 1),
        'others': (0, 2)
    }

    cell_mapping_table24 = {
        'E': (0, 0),
        'C': (0, 1),
        'others': (0, 2)
    }

    cell_mapping_table25 = {
        '範疇或類別': (0, 0),
        '排放源': (0, 1),
        '係數來源': (0, 2),
        '係數名稱': (0, 3),
        '氣體': (0, 4),
        '溫室氣體排放係數': (0, 5),
        '單位': (0, 6)
    }

    cell_mapping_table34 = {
        'B': (0, 0),
        'C': (0, 1),
        'D': (0, 2),
        'E': (0, 3),
        'F': (0, 4),
        'G': (0, 5),
        'H': (0, 6),
        'I': (0, 7),
        'J': (0, 8)
    }

    excel_data_table1 = read_excel_data(excel_path, '表1.基本資料', start_cells)
    fill_word_table(
        word_path=word_path,
        output_path=output_path,
        table_index=0,
        excel_data=excel_data_table1,
        cell_mapping=cell_mapping_table0,
        start_row=1
    )

    excel_data_table2 = read_excel_data(excel_path, '表2.排放源鑑別', start_cells)

    fill_word_table(
        word_path=output_path,
        output_path=output_path,
        table_index=1,
        excel_data=excel_data_table2,
        cell_mapping=cell_mapping_table1,
        start_row=1
    )
    fill_word_table(
        word_path=output_path,
        output_path=output_path,
        table_index=3,
        excel_data=excel_data_table2,
        cell_mapping=cell_mapping_table3,
        start_row=1
    )

    fill_word_table(
        word_path=output_path,
        output_path=output_path,
        table_index=5,
        excel_data=excel_data_table2,
        cell_mapping=cell_mapping_table5,
        start_row=1
    )

    fill_word_table(
        word_path=output_path,
        output_path=output_path,
        table_index=6,
        excel_data=excel_data_table2,
        cell_mapping=cell_mapping_table6,
        start_row=1
    )

    fill_word_table(
        word_path=output_path,
        output_path=output_path,
        table_index=7,
        excel_data=excel_data_table2,
        cell_mapping=cell_mapping_table7,
        start_row=1
    )

    fill_word_table(
        word_path=output_path,
        output_path=output_path,
        table_index=8,
        excel_data=excel_data_table2,
        cell_mapping=cell_mapping_table8,
        start_row=1
    )

    fill_word_table(
        word_path=output_path,
        output_path=output_path,
        table_index=10,
        excel_data=excel_data_table2,
        cell_mapping=cell_mapping_table10,
        start_row=1
    )

    fill_word_table(
        word_path=output_path,
        output_path=output_path,
        table_index=11,
        excel_data=excel_data_table2,
        cell_mapping=cell_mapping_table11,
        start_row=1
    )
    fill_word_table(
        word_path=output_path,
        output_path=output_path,
        table_index=13,
        excel_data=excel_data_table2,
        cell_mapping=cell_mapping_table13,
        start_row=1
    )
    fill_word_table(
        word_path=output_path,
        output_path=output_path,
        table_index=14,
        excel_data=excel_data_table2,
        cell_mapping=cell_mapping_table14,
        start_row=1
    )
    fill_word_table(
        word_path=output_path,
        output_path=output_path,
        table_index=15,
        excel_data=excel_data_table2,
        cell_mapping=cell_mapping_table15,
        start_row=1
    )

    fill_word_table(
        word_path=output_path,
        output_path=output_path,
        table_index=16,
        excel_data=excel_data_table2,
        cell_mapping=cell_mapping_table16,
        start_row=1
    )
    excel_data_table3 = read_excel_data(excel_path, '表3.活動數據', start_cells)
    fill_word_table(
        word_path=output_path,
        output_path=output_path,
        table_index=23,
        excel_data=excel_data_table3,
        cell_mapping=cell_mapping_table23,
        start_row=1
    )

    fill_word_table(
        word_path=output_path,
        output_path=output_path,
        table_index=24,
        excel_data=excel_data_table2,
        cell_mapping=cell_mapping_table24,
        start_row=1
    )

    excel_data_table5 = read_excel_data_pandas(excel_path, '表5.排放係數')
    fill_word_table(
        word_path=output_path,
        output_path=output_path,
        table_index=25,
        excel_data=excel_data_table5,
        cell_mapping=cell_mapping_table25,
        start_row=1
    )
    merge_cells_in_table_25(
    word_path=output_path,
    output_path=output_path,
    table_index=25
    )

    excel_data_table8 = read_excel_data(excel_path, '表8.不確定分析', start_cells)
    fill_word_table(
        word_path=output_path,
        output_path=output_path,
        table_index=34,
        excel_data=excel_data_table8,
        cell_mapping=cell_mapping_table34,
        start_row=1
    )

    replacement_cells_62 = [
        ('Table6.2_D5', 'D5'),
        ('Table6.2_D6', 'D6'),
        ('Table6.2_D7', 'D7'),
        ('Table6.2_D8', 'D8'),
        ('Table6.2_D9', 'D9'),
        ('Table6.2_D10', 'D10'),
        ('Table6.2_D11', 'D11'),
        ('Table6.2_D17', 'D17'),
        ('Table6.2_D18', 'D18'),
        ('Table6.2_D19', 'D19'),
        ('Table6.2_D20', 'D20'),
        ('Table6.2_D21', 'D21'),
        ('Table6.2_D22', 'D22'),
        ('Table6.2_D23', 'D23'),
        ('Table6.2_D24', 'D24'),
        ('Table6.2_D25', 'D25'),
        ('Table6.2_D26', 'D26'),
        ('Table6.2_D27', 'D27'),
        ('Table6.2_D28', 'D28'),
        ('Table6.2_D29', 'D29'),
        ('Table6.2_D30', 'D30'),
        ('Table6.2_D31', 'D31'),
        ('Table6.2_D32', 'D32'),
        ('Table6.2_D33', 'D33')
    ]

    cell_values_62 = read_excel_cells(
        excel_path,
        '表6.2溫室氣體排放量 (範疇1&2, 類別1-15)',
        [cell for _, cell in replacement_cells_62]
    )

    replacements_62 = [(old_text, cell_values_62[cell]) for old_text, cell in replacement_cells_62]

    replace_texts_in_word(
        word_path=output_path,
        output_path=output_path,
        replacements=replacements_62
    )

    replacement_cells_61 = [
        ('Table6.1_J4', 'J4'),
        ('Table6.1_C24', 'C24'),
        ('Table6.1_C25', 'C25'),
        ('Table6.1_G21', 'G21'),
        ('Table6.1_H21', 'H21'),
        ('Table6.1_J21', 'J21'),
        ('Table6.1_K21', 'K21'),
        ('Table6.1_G22', 'G22'),
        ('Table6.1_H23', 'H23'),
        ('Table6.1_C13', 'C13'),
        ('Table6.1_D13', 'D13'),
        ('Table6.1_E13', 'E13'),
        ('Table6.1_F13', 'F13'),
        ('Table6.1_G13', 'G13'),
        ('Table6.1_H13', 'H13'),
        ('Table6.1_I13', 'I13'),
        ('Table6.1_C15', 'C15'),
        ('Table6.1_D15', 'D15'),
        ('Table6.1_E15', 'E15'),
        ('Table6.1_F15', 'F15'),
        ('Table6.1_G15', 'G15'),
        ('Table6.1_H15', 'H15'),
        ('Table6.1_I15', 'I15'),
        ('Table6.1_C21', 'C21'),
        ('Table6.1_D21', 'D21'),
        ('Table6.1_E21', 'E21'),
        ('Table6.1_F21', 'F21'),
        ('Table6.1_C22', 'C22'),
        ('Table6.1_D22', 'D22'),
        ('Table6.1_E22', 'E22'),
        ('Table6.1_F22', 'F22'),
        ('Table6.1_C23', 'C23'),
        ('Table6.1_D23', 'D23'),
        ('Table6.1_E23', 'E23'),
        ('Table6.1_F23', 'F23'),
        ('Table6.1_G23', 'G23'),
        ('Table6.1_H22', 'H22')
    ]

    cell_values_61 = read_excel_cells(
        excel_path,
        '表6.1溫室氣體排放量(範疇1-2)',
        [cell for _, cell in replacement_cells_61]
    )

    replacements_61 = [(old_text, cell_values_61[cell]) for old_text, cell in replacement_cells_61]

    replace_texts_in_word(
        word_path=output_path,
        output_path=output_path,
        replacements=replacements_61
    )

    replacement_cells_7 = [
        ('Table7_O2', 'O2'),
        ('Table7_Q2', 'Q2')
    ]

    cell_values_7 = read_excel_cells(
        excel_path,
        '表7.數據品質分析',
        [cell for _, cell in replacement_cells_7]
    )

    replacements_7 = [(old_text, cell_values_7[cell]) for old_text, cell in replacement_cells_7]

    replace_texts_in_word(
        word_path=output_path,
        output_path=output_path,
        replacements=replacements_7
    )

    replacement_cells_8 = [
        ('Table8_A23', 'A23'),
        ('Table8_C23', 'C23'),
        ('Table8_E23', 'E23')
    ]

    cell_values_8 = read_excel_cells(
        excel_path,
        '表8.不確定分析',
        [cell for _, cell in replacement_cells_8]
    )

    replacements_8 = [(old_text, cell_values_8[cell]) for old_text, cell in replacement_cells_8]

    replace_texts_in_word(
        word_path=output_path,
        output_path=output_path,
        replacements=replacements_8
    )

    replacement_cells_1 = [
        ('rb_version', 'B2'),
        ('rb_published_year', 'D2'),
        ('rb_published_month', 'D3'),
        ('rb_company_name', 'B5'),
        ('rb_company_address', 'B6'),
        ('rb_initiating_year', 'B8'),
        ('rb_base_year', 'B9'),
        ('rb_reporting_year', 'B10'),
        ('rb_reporting_period', 'B11'),
        ('rb_contact_name', 'B12'),
        ('rb_contact_dept', 'B13'),
        ('rb_contact_phone', 'B14'),
        ('rb_contact_email', 'B15')
    ]

    cell_values_1 = read_excel_cells(
        excel_path,
        '表1.基本資料',
        [cell for _, cell in replacement_cells_1]
    )

    replacements_1 = [(old_text, cell_values_1[cell]) for old_text, cell in replacement_cells_1]

    replace_texts_in_word(
        word_path=output_path,
        output_path=output_path,
        replacements=replacements_1
    )

    empty_check_tables = [1, 3, 5, 6, 7, 8, 10, 11, 13, 14, 15]
    insert_if_empty_tables(
        word_path=output_path,
        output_path=output_path,
        table_indices=empty_check_tables
    )

    print(f"Word saved as {output_file_name} at {output_path}")

if __name__ == "__main__":
    # For developer testing only
    with open("test_config.json", "r", encoding="utf-8") as f:
        config = json.load(f)
    main_with_inputs(
        config["excel_path"],
        config["word_path"],
        config["output_folder"],
        config["output_file_name"]
    )
